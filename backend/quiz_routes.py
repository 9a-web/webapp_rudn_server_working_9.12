"""
Раздел «Тесты по лекциям» — генерация карточек-тестов из текста лекции.

Экспортирует factory `create_quiz_router(db)` → APIRouter с префиксом `/quiz`.
Модель: deepseek-v4-flash через шлюз OpenModel (Anthropic Messages формат).

Режимы (mode):
  - multiple_choice — вопрос + 4 варианта, 1 правильный;
  - true_false      — утверждение «Верно/Неверно»;
  - flashcard       — карточка-перевёртыш (вопрос → ответ).

Эндпоинты (все требуют JWT):
  POST   /api/quiz/extract-text      — извлечь текст из файла (.txt/.md/.pdf/.docx)
  POST   /api/quiz/generate          — сгенерировать тест из текста лекции
  GET    /api/quiz/list              — список тестов пользователя (сводка)
  GET    /api/quiz/{quiz_id}         — полный тест с вопросами (только владелец)
  DELETE /api/quiz/{quiz_id}         — удалить тест (только владелец)
  POST   /api/quiz/{quiz_id}/attempt — сохранить попытку прохождения и получить результат
"""

import os
import io
import re
import json
import uuid
import logging
from datetime import datetime, timezone
from typing import List, Optional, Dict, Any

import httpx
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File
from pydantic import BaseModel, Field

from auth_utils import get_current_user_required, check_rate_limit

logger = logging.getLogger(__name__)

# ─── Конфигурация OpenModel / deepseek-v4-flash ───
OPENMODEL_BASE_URL = os.environ.get("OPENMODEL_BASE_URL", "https://api.openmodel.ai/v1").rstrip("/")
OPENMODEL_API_KEY = os.environ.get("OPENMODEL_API_KEY", "")
OPENMODEL_QUIZ_MODEL = os.environ.get("OPENMODEL_QUIZ_MODEL", "deepseek-v4-flash")

MAX_LECTURE_CHARS = 24000   # ограничиваем вход (стоимость/латентность)
MIN_LECTURE_CHARS = 40
GENERATE_RATE_LIMIT = 30    # генераций в час на пользователя
MAX_UPLOAD_BYTES = 12 * 1024 * 1024  # 12 MB
SUPPORTED_MODES = ("multiple_choice", "true_false", "flashcard")


# ─── Pydantic схемы запросов ───
class GenerateQuizRequest(BaseModel):
    text: str = Field(..., min_length=1)
    title: Optional[str] = None
    num_questions: int = Field(10, ge=3, le=20)
    language: Optional[str] = "ru"
    mode: str = "multiple_choice"


class AttemptRequest(BaseModel):
    # Для MCQ/true_false — индекс выбранного варианта (int).
    # Для flashcard — самооценка «знал» (true/1) / «не знал» (false/0).
    answers: List[Optional[Any]] = Field(default_factory=list)


# ─── Утилиты ───
def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _owner_filter(uid: Optional[str], tid: Optional[int]) -> Dict[str, Any]:
    ors: List[Dict[str, Any]] = []
    if uid:
        ors.append({"uid": uid})
    if tid is not None:
        ors.append({"tid": tid})
    if not ors:
        return {"_never_match_": str(uuid.uuid4())}
    return {"$or": ors}


def _build_system_prompt(language: str, mode: str) -> str:
    lang_line = "русском языке" if (language or "ru").startswith("ru") else "том же языке, что и лекция"
    if mode == "flashcard":
        role = "составляет учебные карточки для запоминания по лекциям"
    else:
        role = "составляет качественные тестовые вопросы по лекциям для проверки знаний студентов"
    return (
        f"Ты — опытный преподаватель университета, который {role}. "
        f"Весь текст формулируй на {lang_line}. "
        "Опирайся на ключевые идеи лекции, а не на мелкие детали. "
        "Отвечай ИСКЛЮЧИТЕЛЬНО валидным JSON-объектом без какого-либо текста до или после него."
    )


def _build_user_prompt(text: str, num_questions: int, mode: str) -> str:
    lecture = f"\n\nТекст лекции:\n\"\"\"\n{text}\n\"\"\""
    if mode == "true_false":
        return (
            f"Составь ровно {num_questions} утверждений по тексту лекции ниже. "
            "Примерно половина утверждений должна быть истинной, половина — ложной. "
            "Для каждого дай короткое пояснение (1–2 предложения).\n\n"
            "Верни JSON строго такого вида (без markdown, без ```):\n"
            '{"questions":[{"statement":"...","is_true":true,"explanation":"..."}]}\n\n'
            "Поле is_true — булево: true, если утверждение верно; false, если ложно."
            + lecture
        )
    if mode == "flashcard":
        return (
            f"Составь ровно {num_questions} карточек для запоминания по тексту лекции ниже. "
            "На лицевой стороне (front) — короткий вопрос или термин, "
            "на обратной (back) — краткий, но полный и точный ответ/определение.\n\n"
            "Верни JSON строго такого вида (без markdown, без ```):\n"
            '{"questions":[{"front":"...","back":"..."}]}'
            + lecture
        )
    # multiple_choice (default)
    return (
        f"Составь ровно {num_questions} тестовых вопросов с выбором ответа по тексту лекции ниже.\n\n"
        "Требования к каждому вопросу:\n"
        "- ровно 4 варианта ответа;\n"
        "- ровно один правильный вариант;\n"
        "- короткое пояснение (1–2 предложения), почему правильный ответ верен.\n\n"
        "Верни JSON строго такого вида (без markdown, без ```):\n"
        '{"questions":[{"question":"...","options":["A","B","C","D"],'
        '"correct_index":0,"explanation":"..."}]}\n\n'
        "Поле correct_index — индекс правильного варианта в массиве options (0, 1, 2 или 3)."
        + lecture
    )


def _extract_text_blocks(data: Dict[str, Any]) -> str:
    """Из ответа Anthropic Messages берём только блоки type == 'text'
    (deepseek-v4-flash дополнительно отдаёт блок 'thinking', который нам не нужен)."""
    content = data.get("content") or []
    parts: List[str] = []
    for block in content:
        if isinstance(block, dict) and block.get("type") == "text":
            parts.append(block.get("text") or "")
    return "".join(parts).strip()


def _load_json_questions(raw: str) -> List[Any]:
    """Надёжно достаёт массив questions из ответа модели."""
    if not raw:
        raise ValueError("empty model output")
    candidate = raw.strip()
    if candidate.startswith("```"):
        candidate = re.sub(r"^```[a-zA-Z]*\n?", "", candidate)
        candidate = re.sub(r"\n?```$", "", candidate).strip()

    def _try(s: str):
        try:
            return json.loads(s)
        except json.JSONDecodeError:
            return None

    data = _try(candidate)
    if data is None:
        start = candidate.find("{")
        end = candidate.rfind("}")
        if start != -1 and end != -1 and end > start:
            data = _try(candidate[start:end + 1])
    if data is None:
        raise ValueError("model did not return valid JSON")

    questions = data.get("questions") if isinstance(data, dict) else None
    if not isinstance(questions, list) or not questions:
        raise ValueError("JSON missing non-empty 'questions' array")
    return questions


def _parse_quiz_json(raw: str, mode: str) -> List[Dict[str, Any]]:
    """Валидирует и нормализует вопросы под нужный режим."""
    items = _load_json_questions(raw)
    cleaned: List[Dict[str, Any]] = []

    for item in items:
        if not isinstance(item, dict):
            continue

        if mode == "flashcard":
            front = (item.get("front") or item.get("question") or "").strip()
            back = (item.get("back") or item.get("answer") or "").strip()
            if not front or not back:
                continue
            cleaned.append({
                "id": str(uuid.uuid4()),
                "question": front,
                "answer": back,
                "explanation": str(item.get("explanation") or "").strip(),
            })

        elif mode == "true_false":
            stmt = (item.get("statement") or item.get("question") or "").strip()
            if not stmt:
                continue
            is_true = item.get("is_true")
            if isinstance(is_true, str):
                is_true = is_true.strip().lower() in ("true", "верно", "да", "1", "yes")
            is_true = bool(is_true)
            cleaned.append({
                "id": str(uuid.uuid4()),
                "question": stmt,
                "options": ["Верно", "Неверно"],
                "correct_index": 0 if is_true else 1,
                "explanation": str(item.get("explanation") or "").strip(),
            })

        else:  # multiple_choice
            q_text = (item.get("question") or "").strip()
            options = item.get("options")
            if not q_text or not isinstance(options, list) or len(options) != 4:
                continue
            options = [str(o).strip() for o in options]
            if any(not o for o in options):
                continue
            try:
                ci = int(item.get("correct_index"))
            except (TypeError, ValueError):
                continue
            if ci < 0 or ci > 3:
                continue
            cleaned.append({
                "id": str(uuid.uuid4()),
                "question": q_text,
                "options": options,
                "correct_index": ci,
                "explanation": str(item.get("explanation") or "").strip(),
            })

    if not cleaned:
        raise ValueError("no valid questions after validation")
    return cleaned


# ─── Извлечение текста из файлов ───
def _extract_pdf(raw: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(raw))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _extract_docx(raw: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(raw))
    parts: List[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


async def _call_deepseek(system_prompt: str, user_prompt: str, max_tokens: int = 4096) -> str:
    if not OPENMODEL_API_KEY:
        raise HTTPException(status_code=503, detail="quiz_generation_not_configured")

    headers = {
        "Authorization": f"Bearer {OPENMODEL_API_KEY}",
        "Content-Type": "application/json",
        "anthropic-version": "2023-06-01",
    }
    payload = {
        "model": OPENMODEL_QUIZ_MODEL,
        "max_tokens": max_tokens,
        "temperature": 0.3,
        "system": system_prompt,
        "messages": [{"role": "user", "content": user_prompt}],
    }

    try:
        async with httpx.AsyncClient(timeout=120.0) as client:
            resp = await client.post(f"{OPENMODEL_BASE_URL}/messages", headers=headers, json=payload)
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="Модель не успела ответить. Попробуйте сократить лекцию.")
    except httpx.HTTPError as exc:
        logger.error(f"[quiz] OpenModel network error: {exc}")
        raise HTTPException(status_code=502, detail="Ошибка соединения с сервисом генерации.")

    if resp.status_code != 200:
        logger.error(f"[quiz] OpenModel {resp.status_code}: {resp.text[:400]}")
        raise HTTPException(status_code=502, detail="Сервис генерации вернул ошибку. Попробуйте позже.")

    return _extract_text_blocks(resp.json())


def create_quiz_router(db) -> APIRouter:
    """Создаёт и возвращает APIRouter с эндпоинтами раздела «Тесты»."""
    router = APIRouter(prefix="/quiz", tags=["quiz"])

    def _identity(current_user: Dict[str, Any]):
        uid_ = current_user.get("uid")
        tid_raw = current_user.get("tid")
        try:
            tid_ = int(tid_raw) if tid_raw is not None else None
        except (TypeError, ValueError):
            tid_ = None
        if not uid_ and tid_ is None:
            raise HTTPException(status_code=401, detail="invalid_token_no_identity")
        return uid_, tid_

    def _summary(doc: Dict[str, Any]) -> Dict[str, Any]:
        return {
            "id": doc.get("id"),
            "title": doc.get("title"),
            "mode": doc.get("mode", "multiple_choice"),
            "num_questions": doc.get("num_questions"),
            "language": doc.get("language"),
            "source_preview": doc.get("source_preview"),
            "created_at": doc.get("created_at"),
            "best_score": doc.get("best_score"),
            "best_percent": doc.get("best_percent"),
            "attempts_count": doc.get("attempts_count", 0),
        }

    @router.post("/extract-text")
    async def extract_text(
        file: UploadFile = File(...),
        current_user: Dict[str, Any] = Depends(get_current_user_required),
    ):
        _identity(current_user)
        name = file.filename or ""
        ext = name.lower().rsplit(".", 1)[-1] if "." in name else ""
        raw = await file.read()
        if not raw:
            raise HTTPException(status_code=400, detail="Пустой файл.")
        if len(raw) > MAX_UPLOAD_BYTES:
            raise HTTPException(status_code=413, detail="Файл слишком большой (макс. 12 МБ).")

        try:
            if ext in ("txt", "md", "text"):
                text = raw.decode("utf-8", errors="ignore")
            elif ext == "pdf":
                text = _extract_pdf(raw)
            elif ext == "docx":
                text = _extract_docx(raw)
            elif ext == "doc":
                raise HTTPException(status_code=400, detail="Формат .doc не поддерживается — сохраните как .docx или .txt.")
            else:
                raise HTTPException(status_code=400, detail="Поддерживаются файлы .txt, .md, .pdf, .docx")
        except HTTPException:
            raise
        except Exception as exc:
            logger.error(f"[quiz] extract '{name}' failed: {exc}")
            raise HTTPException(status_code=422, detail="Не удалось извлечь текст из файла.")

        text = (text or "").strip()
        if len(text) < MIN_LECTURE_CHARS:
            raise HTTPException(status_code=422, detail="В файле слишком мало текста (возможно, это скан без распознавания).")

        truncated = len(text) > MAX_LECTURE_CHARS
        return {
            "text": text[:MAX_LECTURE_CHARS],
            "filename": name,
            "chars": len(text),
            "truncated": truncated,
        }

    @router.post("/generate")
    async def generate_quiz(
        body: GenerateQuizRequest,
        current_user: Dict[str, Any] = Depends(get_current_user_required),
    ):
        uid_, tid_ = _identity(current_user)

        mode = body.mode if body.mode in SUPPORTED_MODES else "multiple_choice"

        rl_key = str(uid_ or tid_)
        if not check_rate_limit(rl_key, "quiz_generate", GENERATE_RATE_LIMIT, 3600):
            raise HTTPException(status_code=429, detail="Слишком много генераций. Попробуйте позже.")

        text = (body.text or "").strip()
        if len(text) < MIN_LECTURE_CHARS:
            raise HTTPException(status_code=400, detail="Текст лекции слишком короткий для генерации теста.")
        truncated = len(text) > MAX_LECTURE_CHARS
        if truncated:
            text = text[:MAX_LECTURE_CHARS]

        system_prompt = _build_system_prompt(body.language or "ru", mode)
        user_prompt = _build_user_prompt(text, body.num_questions, mode)
        max_tokens = min(8000, 600 + body.num_questions * 320)

        raw = await _call_deepseek(system_prompt, user_prompt, max_tokens=max_tokens)
        try:
            questions = _parse_quiz_json(raw, mode)
        except ValueError as exc:
            logger.error(f"[quiz] parse error ({mode}): {exc}; raw[:300]={raw[:300]!r}")
            raise HTTPException(status_code=502, detail="Не удалось разобрать ответ модели. Попробуйте ещё раз.")

        default_title = "Карточки по лекции" if mode == "flashcard" else "Тест по лекции"
        title = (body.title or "").strip() or default_title
        quiz_doc = {
            "id": str(uuid.uuid4()),
            "uid": uid_,
            "tid": tid_,
            "title": title[:120],
            "mode": mode,
            "language": body.language or "ru",
            "source_preview": text[:280],
            "source_length": len(text),
            "source_truncated": truncated,
            "num_questions": len(questions),
            "questions": questions,
            "model": OPENMODEL_QUIZ_MODEL,
            "created_at": _now_iso(),
            "best_score": None,
            "best_percent": None,
            "attempts_count": 0,
        }
        await db.quizzes.insert_one(quiz_doc)
        quiz_doc.pop("_id", None)
        logger.info(f"[quiz] generated quiz id={quiz_doc['id']} mode={mode} q={len(questions)} uid={uid_}")
        return quiz_doc

    @router.get("/list")
    async def list_quizzes(current_user: Dict[str, Any] = Depends(get_current_user_required)):
        uid_, tid_ = _identity(current_user)
        cursor = db.quizzes.find(_owner_filter(uid_, tid_), {"_id": 0, "questions": 0}).sort("created_at", -1)
        docs = await cursor.to_list(length=200)
        return {"quizzes": [_summary(d) for d in docs], "total": len(docs)}

    @router.get("/{quiz_id}")
    async def get_quiz(quiz_id: str, current_user: Dict[str, Any] = Depends(get_current_user_required)):
        uid_, tid_ = _identity(current_user)
        flt = {"id": quiz_id, **_owner_filter(uid_, tid_)}
        doc = await db.quizzes.find_one(flt, {"_id": 0})
        if not doc:
            raise HTTPException(status_code=404, detail="Тест не найден")
        return doc

    @router.delete("/{quiz_id}")
    async def delete_quiz(quiz_id: str, current_user: Dict[str, Any] = Depends(get_current_user_required)):
        uid_, tid_ = _identity(current_user)
        flt = {"id": quiz_id, **_owner_filter(uid_, tid_)}
        res = await db.quizzes.delete_one(flt)
        if res.deleted_count == 0:
            raise HTTPException(status_code=404, detail="Тест не найден")
        await db.quiz_attempts.delete_many({"quiz_id": quiz_id})
        return {"status": "ok", "deleted": quiz_id}

    @router.post("/{quiz_id}/attempt")
    async def submit_attempt(
        quiz_id: str,
        body: AttemptRequest,
        current_user: Dict[str, Any] = Depends(get_current_user_required),
    ):
        uid_, tid_ = _identity(current_user)
        flt = {"id": quiz_id, **_owner_filter(uid_, tid_)}
        quiz = await db.quizzes.find_one(flt, {"_id": 0})
        if not quiz:
            raise HTTPException(status_code=404, detail="Тест не найден")

        mode = quiz.get("mode", "multiple_choice")
        questions = quiz.get("questions") or []
        answers = body.answers or []
        results: List[Dict[str, Any]] = []
        correct = 0

        for i, q in enumerate(questions):
            raw_ans = answers[i] if i < len(answers) else None
            if mode == "flashcard":
                # самооценка: знал = true/1
                knew = bool(raw_ans) if not isinstance(raw_ans, str) else raw_ans.strip().lower() in ("true", "1", "yes", "да")
                if knew:
                    correct += 1
                results.append({"question_id": q.get("id"), "knew": knew})
            else:
                try:
                    your = int(raw_ans) if raw_ans is not None else None
                except (TypeError, ValueError):
                    your = None
                is_correct = your is not None and your == q.get("correct_index")
                if is_correct:
                    correct += 1
                results.append({
                    "question_id": q.get("id"),
                    "correct_index": q.get("correct_index"),
                    "your_index": your,
                    "is_correct": is_correct,
                })

        total = len(questions)
        percent = round((correct / total) * 100) if total else 0

        attempt_doc = {
            "id": str(uuid.uuid4()),
            "quiz_id": quiz_id,
            "uid": uid_,
            "tid": tid_,
            "mode": mode,
            "score": correct,
            "total": total,
            "percent": percent,
            "created_at": _now_iso(),
        }
        await db.quiz_attempts.insert_one(attempt_doc)

        prev_best = quiz.get("best_score")
        update: Dict[str, Any] = {"$inc": {"attempts_count": 1}}
        if prev_best is None or correct > prev_best:
            update["$set"] = {"best_score": correct, "best_percent": percent}
        await db.quizzes.update_one({"id": quiz_id}, update)

        return {
            "score": correct,
            "total": total,
            "percent": percent,
            "results": results,
        }

    return router
